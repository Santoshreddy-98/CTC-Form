from flask import Flask, render_template, request, jsonify, send_file, make_response
from docx import Document
import os

app = Flask(__name__)

# Making docs file

def generate_document(data):
    template_file_path = './Annexure - Annually and monthly.docx'
    output_file_path = data["name"]
    print(output_file_path)
    variables = {
        "${EMPLOYEE_NAME}": data["name"],
        "${EMPLOYEE_DESIGNATION}": data['designation'],
        "${EMP_LOCATION}": data['location'],
        "${SALARY}":data['salary'],
        "${MON_SALARY}":data['monthly_salary'],
        "${DOJ}": data['DOJ'],
        "${ANNUAL_BASIC_CTC}": data['basic_percentage'],
        "${MON_BASIC_CTC}": data['monthly_basic'],
        "${ANNUAL_HRA}": data['hra_percentage'],
        "${MONTHLY_HRA}": data['monthly_hra'],
        "${SPL_ALLOWANCE}": data['special_allowance_percentage'],
        "${MON_ALLOWANCE}": data['monthly_special_allowance'],
        "${A_CONVEYANCE}": data['conveyance_percentage'],
        "${MON_CONVEYANCE}": data['monthly_conveyance'],
        "${ANNUAL_TOTAL}": data['annual_ctc'],
        "${MONTHLY_TOTAL}": data['monthly_ctc'],
        "${ANL_VAR_PAY}": data["variablePay"],
        "${MON_VAR_PAY}": data["monthly_variable_pay"],
        "${CCPF}": data["CCPF"],
        "${MON_CCPF}":data["monthly_CCPF"],
        "${ANNUAL_ESIC_SHARE}": "-",
        "${MONTHLY_ESIC}": "-",
        "${ANL_COM_CON_PF}": "-",
        "${MON_COM_CON_PF}": "-",
        "${ANNUAL_DEDUCTIONS}": "-",
        "${MONTHLY_DEDUCTIONS}": "-",
        "${EMPLOYER_SHARE_PF}": "-",
        "${MON_EMPLOYER_SHARE_PF}": "-",
        "${ANL_GROSS_PAY}": data['salary'],
        "${MON_GROSS_PAY}": "-",
        "${ANL_ESIC_EMPLOY_SHR}": "-",
        "${MON_ESIC_EMPLOY_SHAR}": "-",
        "${ANL_EMPLOY_SHR_PF}": "-",
        "${MON_EMPLOY_SHAR_PF}": "-",
        "${ANL_MED_INS}": "-",
        "${MON_MED_INS}": "-",
        "${ANL_PRO_TAX}": "-",
        "${MON_PRO_TAX}": "-",
        "${ANL_NET_PAY}": "-",
        "${MON_NET_PAY}": "-",
        "${CTC_EMPLOYR_PF_ESIC_SHARE}": "-",
    }

    template_document = load_template(template_file_path)
    for placeholder, replacement_value in variables.items():
        replace_text_in_paragraphs(
            template_document.paragraphs, placeholder, replacement_value)
        replace_text_in_tables(template_document.tables,
                               placeholder, replacement_value)

    save_document(template_document, output_file_path)

    return output_file_path


def load_template(file_path):
    return Document(file_path)


def replace_text_in_paragraphs(paragraphs, placeholder, replacement_value):
    for paragraph in paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(
                placeholder, replacement_value)


def replace_text_in_tables(tables, placeholder, replacement_value):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(
                    cell.paragraphs, placeholder, replacement_value)


def save_document(document, file_path):
    document.save(file_path)


@app.route('/')
def index():
    port = 7393  # Change the port number to your desired port
    return render_template('form.html', port=port)


@app.route('/calculate', methods=['POST'])
def calculate():
    data = request.get_json()
    file_path = generate_document(data)
    print(data['name'])

    filename = data["name"]

    response = make_response(send_file(file_path))

    response.headers["Content-Disposition"] = f"attachment; filename={filename}"
    print(response)

    return response


if __name__ == '__main__':
    port = 7393  # Change the port number to your desired port
    app.run(debug=True, port=port)
